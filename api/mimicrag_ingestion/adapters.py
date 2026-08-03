from __future__ import annotations

import email
import hashlib
import html
import io
import json
import mimetypes
import re
from dataclasses import dataclass, field
from email import policy
from pathlib import Path
from typing import Any, Callable


@dataclass
class StructuredBlock:
    type: str
    text: str
    page: int = 0
    section_path: str = ""
    coordinates: tuple[float, float, float, float] | None = None
    headers: list[str] = field(default_factory=list)
    rows: list[list[str]] = field(default_factory=list)
    caption: str = ""
    references: list[str] = field(default_factory=list)


@dataclass
class NormalizedDocument:
    source_uri: str
    format: str
    text: str
    title: str = ""
    blocks: list[StructuredBlock] = field(default_factory=list)
    metadata: dict[str, Any] = field(default_factory=dict)
    warnings: list[str] = field(default_factory=list)

    @property
    def content_hash(self) -> str:
        normalized = re.sub(r"\s+", " ", self.text).strip().encode()
        return hashlib.sha256(normalized).hexdigest()

    def ingestion_metadata(self) -> dict[str, Any]:
        blocks = []
        for block in self.blocks:
            item = {"type": block.type, "page": block.page, "section_path": block.section_path}
            if block.coordinates is not None:
                item["coordinates"] = list(block.coordinates)
            if block.headers:
                item["headers"] = block.headers
            if block.rows:
                item["rows"] = block.rows
            if block.caption:
                item["caption"] = block.caption
            if block.references:
                item["references"] = block.references
            blocks.append(item)
        return {**self.metadata, "source_format": self.format, "content_hash": self.content_hash,
                "structured_content": blocks}


def _decode(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-16", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            pass
    return data.decode("utf-8", errors="replace")


def _markdown_blocks(text: str) -> list[StructuredBlock]:
    blocks: list[StructuredBlock] = []
    sections: list[str] = []
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        heading = re.match(r"^\s{0,3}(#{1,6})\s+(.+)$", line)
        if heading:
            level, label = len(heading[1]), heading[2].strip()
            sections = sections[: level - 1] + [label]
            blocks.append(StructuredBlock("heading", label, section_path=" > ".join(sections)))
        elif "|" in line and i + 1 < len(lines) and re.match(r"^\s*\|?\s*:?-+", lines[i + 1]):
            headers = [cell.strip() for cell in line.strip(" |").split("|")]
            rows: list[list[str]] = []
            i += 2
            while i < len(lines) and "|" in lines[i]:
                rows.append([cell.strip() for cell in lines[i].strip(" |").split("|")])
                i += 1
            rendered = " | ".join(headers) + "\n" + "\n".join(" | ".join(row) for row in rows)
            blocks.append(StructuredBlock("table", rendered, section_path=" > ".join(sections), headers=headers, rows=rows))
            continue
        elif line.strip():
            kind = "code" if line.lstrip().startswith("```") else "paragraph"
            refs = re.findall(r"\[[^]]+\]\(([^)]+)\)", line)
            blocks.append(StructuredBlock(kind, line, section_path=" > ".join(sections), references=refs))
        i += 1
    return blocks


def _html_document(data: bytes, uri: str, title: str) -> NormalizedDocument:
    raw = _decode(data)
    clean = re.sub(r"(?is)<(script|style).*?>.*?</\1>", " ", raw)
    found_title = re.search(r"(?is)<title[^>]*>(.*?)</title>", clean)
    title = title or (html.unescape(re.sub(r"<[^>]+>", "", found_title[1])).strip() if found_title else "")
    blocks: list[StructuredBlock] = []
    sections: list[str] = []
    pattern = re.compile(r"(?is)<(h[1-6]|p|li|pre|caption|figcaption|table)\b[^>]*>(.*?)</\1>")
    for match in pattern.finditer(clean):
        tag = match[1].lower()
        value = html.unescape(re.sub(r"(?s)<[^>]+>", " ", match[2]))
        value = re.sub(r"\s+", " ", value).strip()
        if not value:
            continue
        if tag.startswith("h"):
            level = int(tag[1]); sections = sections[: level - 1] + [value]; kind = "heading"
        elif tag == "table":
            rows = [[html.unescape(re.sub(r"<[^>]+>", "", cell)).strip()
                     for cell in re.findall(r"(?is)<t[dh][^>]*>(.*?)</t[dh]>", row)]
                    for row in re.findall(r"(?is)<tr[^>]*>(.*?)</tr>", match[2])]
            rows = [row for row in rows if row]
            headers = rows.pop(0) if rows else []
            blocks.append(StructuredBlock("table", value, section_path=" > ".join(sections), headers=headers, rows=rows))
            continue
        elif tag in ("caption", "figcaption"): kind = "caption"
        elif tag == "pre": kind = "code"
        else: kind = "paragraph"
        blocks.append(StructuredBlock(kind, value, section_path=" > ".join(sections)))
    text = "\n\n".join(block.text for block in blocks) or html.unescape(re.sub(r"(?s)<[^>]+>", " ", clean))
    return NormalizedDocument(uri, "html", text, title, blocks)


def _email_document(data: bytes, uri: str, title: str) -> NormalizedDocument:
    message = email.message_from_bytes(data, policy=policy.default)
    parts: list[str] = []
    for part in message.walk():
        if part.get_content_disposition() == "attachment":
            continue
        if part.get_content_type() in ("text/plain", "text/html"):
            content = part.get_content()
            parts.append(re.sub(r"(?s)<[^>]+>", " ", content) if part.get_content_type() == "text/html" else content)
    text = "\n\n".join(parts)
    metadata = {"email": {key.lower(): str(message.get(key, "")) for key in ("From", "To", "Cc", "Date", "Message-ID")}}
    return NormalizedDocument(uri, "email", text, title or str(message.get("Subject", "")), _markdown_blocks(text), metadata)


def _json_document(data: bytes, uri: str, title: str) -> NormalizedDocument:
    value = json.loads(_decode(data))
    text = json.dumps(value, ensure_ascii=False, indent=2)
    blocks: list[StructuredBlock] = []
    if isinstance(value, list) and value and all(isinstance(row, dict) for row in value):
        headers = list(dict.fromkeys(key for row in value for key in row))
        rows = [[str(row.get(key, "")) for key in headers] for row in value]
        blocks.append(StructuredBlock("table", text, headers=headers, rows=rows))
    else:
        blocks.append(StructuredBlock("data", text))
    return NormalizedDocument(uri, "json", text, title, blocks)


def _docx_document(data: bytes, uri: str, title: str) -> NormalizedDocument:
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:
        raise RuntimeError("DOCX ingestion requires the optional 'python-docx' package") from exc
    document = Document(io.BytesIO(data))
    blocks = [StructuredBlock("paragraph", p.text) for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        headers = rows.pop(0) if rows else []
        blocks.append(StructuredBlock("table", "\n".join(" | ".join(r) for r in [headers, *rows]), headers=headers, rows=rows))
    text = "\n\n".join(block.text for block in blocks)
    return NormalizedDocument(uri, "docx", text, title, blocks)


def _pdf_document(data: bytes, uri: str, title: str, ocr: Callable[[bytes, int], list[dict[str, Any]]] | None) -> NormalizedDocument:
    try:
        from pypdf import PdfReader  # type: ignore
    except ImportError as exc:
        raise RuntimeError("PDF ingestion requires the optional 'pypdf' package") from exc
    reader = PdfReader(io.BytesIO(data)); blocks: list[StructuredBlock] = []; warnings: list[str] = []
    for page_number, page in enumerate(reader.pages, 1):
        text = page.extract_text() or ""
        if text.strip():
            blocks.append(StructuredBlock("page", text, page=page_number))
        elif ocr:
            for item in ocr(data, page_number):
                blocks.append(StructuredBlock("ocr", str(item["text"]), page_number,
                    coordinates=tuple(item["coordinates"])))  # type: ignore[arg-type]
        else:
            warnings.append(f"page {page_number} has no extractable text; configure OCR")
    return NormalizedDocument(uri, "pdf", "\n\n".join(b.text for b in blocks), title, blocks, warnings=warnings)


class AdapterRegistry:
    """Dependency-light adapter registry. PDF/DOCX parsers are loaded only when used."""

    CODE_EXTENSIONS = {".c", ".cc", ".cpp", ".h", ".hpp", ".py", ".js", ".jsx", ".ts", ".tsx", ".go", ".rs", ".java", ".rb", ".php", ".sh", ".sql"}

    def __init__(self, ocr: Callable[[bytes, int], list[dict[str, Any]]] | None = None):
        self.ocr = ocr

    def parse(self, data: bytes, source_uri: str, *, media_type: str = "", title: str = "") -> NormalizedDocument:
        suffix = Path(source_uri.split("?", 1)[0]).suffix.lower()
        media_type = (media_type or mimetypes.guess_type(source_uri)[0] or "").split(";", 1)[0]
        if suffix == ".pdf" or media_type == "application/pdf": return _pdf_document(data, source_uri, title, self.ocr)
        if suffix == ".docx" or media_type.endswith("wordprocessingml.document"): return _docx_document(data, source_uri, title)
        if suffix in (".html", ".htm") or media_type == "text/html": return _html_document(data, source_uri, title)
        if suffix == ".eml" or media_type == "message/rfc822": return _email_document(data, source_uri, title)
        if suffix in (".json", ".jsonl") or media_type == "application/json": return _json_document(data, source_uri, title)
        text = _decode(data)
        if suffix in (".md", ".markdown") or media_type == "text/markdown": fmt = "markdown"
        elif suffix in self.CODE_EXTENSIONS: fmt = "source_code"
        else: fmt = "text"
        return NormalizedDocument(source_uri, fmt, text, title, _markdown_blocks(text))
